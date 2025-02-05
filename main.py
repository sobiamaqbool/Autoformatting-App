import os
import spacy
from flask import Flask, request, render_template, send_file
from PyPDF2 import PdfReader
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import RGBColor
from docx.shared import Inches
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml

app = Flask(__name__)

UPLOAD_FOLDER = 'uploads'
TEMPLATES_FOLDER = 'templates'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['TEMPLATES_FOLDER'] = TEMPLATES_FOLDER


def is_heading(paragraph):
    if paragraph.style.name.startswith('Heading'):
        return True
    if len(paragraph.text) < 40:
        return True
    return False


def read_docx(file_path):
    doc = Document(file_path)
    data = []
    current_heading = None
    toc_headings = []
    toc_headings_started = False

    for paragraph in doc.paragraphs:
        trimmed_text = paragraph.text.strip()
        if not trimmed_text or len(trimmed_text.split()) < 2:
            continue

        if is_heading(paragraph):
            if current_heading is not None:
                data.append(current_heading)
            current_heading = {'heading': paragraph.text, 'paragraphs': []}
            if toc_headings_started:
                toc_headings.append(paragraph.text)
        elif current_heading is not None:
            current_heading['paragraphs'].append(paragraph.text)

        if not toc_headings_started and paragraph.text == "Lahore College for Women University":
            toc_headings_started = True

    if current_heading is not None:
        data.append(current_heading)

    return data, toc_headings


def clear_footer(footer):
    for paragraph in footer.paragraphs:
        for run in paragraph.runs:
            run.clear()


def create_doc(data, toc_headings, font_name='Times New Roman', font_size=12, line_spacing=1.0):
    doc = Document()

    sections = doc.sections
    for section in sections:
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

        footer = sections[0].footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        fld_xml = f'<w:fldSimple {nsdecls("w")} w:instr="PAGE"/>'
        footer_para._element.append(parse_xml(fld_xml))

        for section in sections[1:]:
            section.footer.is_linked_to_previous = False
            footer = section.footer
            footer.remove()

    font_style_normal = doc.styles['Normal']
    font_style_heading = doc.styles['Heading 1']

    font_style_normal.font.name = font_name
    font_style_normal.font.size = Pt(font_size)

    font_style_heading.font.bold = True

    toc_headings_started = False

    for item in data:
        heading = doc.add_heading(item['heading'], level=1)
        run = heading.runs[0]
        run.font.color.rgb = RGBColor(0, 0, 0)

        if item['heading'] == "Lahore College for Women University":
            doc.add_page_break()

            toc_heading = doc.add_heading('Table of Contents', level=1)
            toc_run = toc_heading.runs[0]
            toc_run.font.color.rgb = RGBColor(0, 0, 0)
            for heading_text in toc_headings:
                toc_heading = doc.add_paragraph(heading_text, style='TOC Heading')
                if toc_heading.text in toc_headings:
                    run = toc_heading.runs[0]
                    run.font.color.rgb = RGBColor(0, 0, 0)

            doc.add_page_break()
            continue

        if not toc_headings_started and item['heading'] == "Lahore College for Women University":
            toc_headings_started = True

        for paragraph_text in item['paragraphs']:
            paragraph = doc.add_paragraph(paragraph_text)
            paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
            paragraph.paragraph_format.line_spacing = line_spacing
            run = paragraph.runs[0]
            run.font.name = font_name
            run.font.size = Pt(font_size)

    output_file_path = os.path.join(app.config['UPLOAD_FOLDER'], 'output.docx')
    doc.save(output_file_path)
    print("Document created successfully!")
    return output_file_path


def extract_text_from_pdf(file):
    pdf_reader = PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text()
    return text


@app.route('/')
def upload_form():
    templates = os.listdir(app.config['TEMPLATES_FOLDER'])
    font_styles = ['Times New Roman', 'Arial', 'Calibri', 'Cambria']
    font_sizes = [8, 9, 10, 11, 12, 14, 16, 18]
    line_spacings = [1.0, 1.15, 1.5, 2.0]
    return render_template('upload.html', templates=templates, font_styles=font_styles, font_sizes=font_sizes, line_spacings=line_spacings)


@app.route('/about')
def about():
    return render_template('about.html')


@app.route('/contactus')
def contactus():
    return render_template('contactus.html')


@app.route('/template')
def template():
    return render_template('template.html')


@app.route('/home')
def home():
    templates = os.listdir(app.config['TEMPLATES_FOLDER'])
    font_styles = ['Times New Roman', 'Arial', 'Calibri', 'Cambria']
    font_sizes = [8, 9, 10, 11, 12, 14, 16, 18]
    line_spacings = [1.0, 1.15, 1.5, 2.0]
    return render_template('upload.html', templates=templates, font_styles=font_styles, font_sizes=font_sizes, line_spacings=line_spacings)


@app.route('/upload', methods=['GET', 'POST'])
def upload_file():
    if request.method == 'POST':
        functionality = request.form.get('functionality')
        uploaded_files = request.files.getlist('file')
        selected_font_style = request.form.get('font_style')
        selected_font_size = int(request.form.get('font_size'))
        selected_line_spacing = float(request.form.get('line_spacing'))
        output_files = []

        for uploaded_file in uploaded_files:
            if uploaded_file.filename != '':
                file_path = os.path.join(app.config['UPLOAD_FOLDER'], uploaded_file.filename)
                uploaded_file.save(file_path)

                if functionality == "extract_text":
                    text = extract_text_from_pdf(file_path)
                    # Handle extracted text
                elif functionality == "create_docx":
                    headings_with_paragraphs, toc_headings = read_docx(file_path)
                    output_file_path = create_doc(headings_with_paragraphs, toc_headings, font_name=selected_font_style,
                                                  font_size=selected_font_size, line_spacing=selected_line_spacing)
                    output_files.append(output_file_path)
                else:
                    pass  # Handle invalid functionality

        if output_files:
            return send_file(output_files[0], as_attachment=True)

    templates = os.listdir(app.config['TEMPLATES_FOLDER'])
    return render_template('upload.html', templates=templates)


if __name__ == '__main__':
    app.run(debug=True)
